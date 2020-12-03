from selenium import webdriver
from selenium.webdriver.chrome.options import Options
import time
import re
import requests
import os

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

def get_picture(dir_path, imgs):
	index = 1
	for url in imgs:
		response = requests.get(url)

		with open(dir_path + str(index) + ".png", "wb") as f:
			f.write(response.content)

		index += 1


def write_doc(driver, document, head_index, imgs, title, price, infos):
	dir_path = "./" + str(head_index) + "/"
	if not os.path.exists("./" + str(head_index)):
		os.makedirs(dir_path)

	# 获得info页的截图
	driver.get_screenshot_as_file(dir_path + "info.png")

	# 添加一级标题
	document.add_heading(str(head_index) + '、', level=0)

	p = document.add_paragraph("");
	p.add_run(title)

	p1 = document.add_paragraph("");
	p1.add_run(price).blod = True

	p1.add_run("\n概要：")
	p1.add_run(infos[0])
	p1.add_run("\n详情：")
	p1.add_run(infos[1])

	#添加info的image
	p1 = document.add_picture(dir_path + "info.png")
	# 按比例缩小
	p1.height = int(document.inline_shapes[0].height * 0.8281573498964804 * 0.31)
	p1.width = int(document.inline_shapes[0].width * 0.6256517205422315 * 0.41)

	get_picture(dir_path, imgs)
	p2 = document.add_paragraph("");

	for i in range(1, len(imgs) + 1):
		p2 = document.add_picture(dir_path + str(i) + ".png")
		p2.height = int(document.inline_shapes[0].height * 0.8281573498964804 * 1.31)
		p2.width = int(document.inline_shapes[0].width * 0.6256517205422315 * 1.51)


def work(urls):
	# docx操作
	document = Document()
	# 设置字体
	document.styles['Normal'].font.name = '宋体'
	document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

	#
	head_index = 1

	# selenium操作
	chrome_options = Options()

	# 注意这个路径需要时可执行路径（chmod 777 dir or 755 dir）
	path = "chromedriver.exe"

	driver = webdriver.Chrome(executable_path=path, chrome_options=chrome_options)

	# 窗口最大化
	driver.maximize_window()

	for url in urls:
		driver.get(url)
		time.sleep(3)

		image_button = driver.find_elements_by_xpath('//button[@class="_90nllgc"]')[0]


		# 打开图片页
		driver.execute_script("arguments[0].click();", image_button)

		# 打开详情页
		info_button = driver.find_elements_by_xpath('//button[@class="_14h2grbm"]')[0]
		driver.execute_script("arguments[0].click();", info_button)

		html = driver.page_source

		# 5张图片
		imgs_src = re.findall('src="(https://z1.muscache.cn/im/pictures/.*?im_w=\d+)"', html)
		title = re.findall('title>(.*?)</title>', html)[0]
		price = re.findall('data-plugin-in-point-id="BOOK_IT_CHINA".*?class="_krjbj".*?(￥.*?)</span>', html)
		info = re.findall("概要(.*?)详情(.*)", driver.page_source)[0]

		infos = []

		for i in info:
			i = i.replace("\n", "</br>")
			i = re.sub(r'<[^>]+>', '', i)
			infos.append(i)

		imgs = []
		for i in imgs_src:
			i = i[: i.find('im_w') - 1]
			imgs.append(i)

		write_doc(driver, document, head_index, imgs, title, price, infos)

		head_index += 1

	# 关闭浏览器
	driver.close()
	driver.quit()
	document.save("airbnb.docx")



urls = [
		"", # 想展示的网站
		]

if __name__ == '__main__':
	work(urls)
